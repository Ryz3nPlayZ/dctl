from __future__ import annotations

from collections.abc import Iterable
from copy import deepcopy
import difflib
import json
from pathlib import Path
import re
import shutil
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from dctl.errors import DctlError


QUESTION_PATTERN = re.compile(r"^\s*(?:\(?[0-9]+[.)]|[A-Za-z][.)]|[ivxlcdm]+[.)])?\s*.+(?:\?|_{3,}|:{1}\s*)\s*$", re.IGNORECASE)
BACKUP_DIRNAME = ".dctl-backups"


def _path(path: str) -> Path:
    resolved = Path(path).expanduser().resolve()
    if not resolved.exists():
        raise DctlError("ELEMENT_NOT_FOUND", f"Document not found: {resolved}")
    return resolved


def _normalize_text(text: str) -> str:
    return " ".join(text.split()).strip().casefold()


def _display_text(text: str) -> str:
    return " ".join(text.split()).strip()


def _parse_json_input(value: str) -> Any:
    candidate = Path(value).expanduser()
    raw = candidate.read_text(encoding="utf-8") if candidate.exists() else value
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        raise DctlError("INVALID_SELECTOR", f"Expected JSON or a JSON file path: {value}") from exc


def _backup_path(path: Path) -> Path:
    backup_dir = path.parent / BACKUP_DIRNAME
    backup_dir.mkdir(exist_ok=True)
    counter = 1
    while True:
        candidate = backup_dir / f"{path.stem}.backup-{counter}{path.suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


def backup(path: str) -> dict[str, Any]:
    doc_path = _path(path)
    destination = _backup_path(doc_path)
    shutil.copy2(doc_path, destination)
    return {"path": str(doc_path), "backup_path": str(destination)}


def _iter_paragraphs(document: Any) -> Iterable[Paragraph]:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _iter_block_items(parent: DocumentObject | _Cell) -> Iterable[Paragraph | Table]:
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise DctlError("INVALID_SELECTOR", "Unsupported DOCX container type.")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _is_question_text(text: str) -> bool:
    stripped = _display_text(text)
    if not stripped:
        return False
    if stripped.endswith("?"):
        return True
    if "____" in stripped:
        return True
    return bool(QUESTION_PATTERN.match(stripped))


def _is_question_paragraph(paragraph: Paragraph) -> bool:
    return _is_question_text(paragraph.text)


def _is_answer_candidate(paragraph: Paragraph) -> bool:
    text = _display_text(paragraph.text)
    if not text:
        return False
    if _is_question_text(text):
        return False
    style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
    if "heading" in style_name or "title" in style_name:
        return False
    return True


def _paragraph_descriptor(paragraph: Paragraph, index: int, kind: str | None = None) -> dict[str, Any]:
    return {
        "index": index,
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "kind": kind or ("question" if _is_question_paragraph(paragraph) else "paragraph"),
        "bold": any(bool(run.bold) for run in paragraph.runs if run.text.strip()),
    }


def inspect(path: str) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    return {
        "path": str(doc_path),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "title": document.core_properties.title or None,
        "subject": document.core_properties.subject or None,
    }


def read(path: str, include_empty: bool = False) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    if not include_empty:
        paragraphs = [paragraph for paragraph in paragraphs if paragraph.strip()]
    return {
        "path": str(doc_path),
        "paragraphs": paragraphs,
        "text": "\n".join(paragraphs),
    }


def paragraphs(path: str, include_empty: bool = True) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    items = []
    for index, paragraph in enumerate(document.paragraphs):
        if not include_empty and not paragraph.text.strip():
            continue
        items.append(_paragraph_descriptor(paragraph, index))
    return {"path": str(doc_path), "items": items}


def _table_title_from_blocks(blocks: list[Paragraph | Table], table_index: int) -> str | None:
    seen_tables = -1
    for position, block in enumerate(blocks):
        if not isinstance(block, Table):
            continue
        seen_tables += 1
        if seen_tables != table_index:
            continue
        for previous in range(position - 1, -1, -1):
            prior = blocks[previous]
            if isinstance(prior, Paragraph) and prior.text.strip():
                return _display_text(prior.text)
        return None
    return None


def _table_matrix(table: Table) -> list[list[str]]:
    return [[_display_text(cell.text) for cell in row.cells] for row in table.rows]


def _table_descriptor(table: Table, index: int, title: str | None = None) -> dict[str, Any]:
    matrix = _table_matrix(table)
    header_row = matrix[0] if matrix else []
    row_labels = [row[0] for row in matrix[1:] if row and row[0]]
    return {
        "index": index,
        "title": title,
        "rows": len(matrix),
        "columns": max((len(row) for row in matrix), default=0),
        "header_row": header_row,
        "row_labels": row_labels,
        "cells": matrix,
    }


def _top_level_blocks(document: Any) -> list[Paragraph | Table]:
    return list(_iter_block_items(document))


def _next_answer_block(document: Any, question_paragraph: Paragraph) -> Paragraph | None:
    blocks = _top_level_blocks(document)
    for index, block in enumerate(blocks):
        if not isinstance(block, Paragraph):
            continue
        if block._p is not question_paragraph._p:
            continue
        next_index = index + 1
        if next_index >= len(blocks):
            return None
        next_block = blocks[next_index]
        if isinstance(next_block, Table):
            return None
        if next_index + 1 < len(blocks) and isinstance(blocks[next_index + 1], Table):
            return None
        return next_block if _is_answer_candidate(next_block) else None
    return None


def worksheet_map(path: str) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    blocks = list(_iter_block_items(document))

    paragraph_items: list[dict[str, Any]] = []
    questions: list[dict[str, Any]] = []
    top_level_index = 0
    for block in blocks:
        if not isinstance(block, Paragraph):
            continue
        descriptor = _paragraph_descriptor(block, top_level_index)
        if descriptor["kind"] == "question":
            next_answer = _next_answer_block(document, block)
            descriptor["existing_answer"] = next_answer.text if next_answer is not None else None
            questions.append(descriptor)
        paragraph_items.append(descriptor)
        top_level_index += 1

    tables = [
        _table_descriptor(table, index, _table_title_from_blocks(blocks, index))
        for index, table in enumerate(document.tables)
    ]
    return {"path": str(doc_path), "paragraphs": paragraph_items, "questions": questions, "tables": tables}


def append(path: str, text_value: str, style: str | None = None) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    backup_result = backup(str(doc_path))
    paragraph = document.add_paragraph(text_value)
    if style:
        paragraph.style = style
    document.save(doc_path)
    return {
        "path": str(doc_path),
        "appended": text_value,
        "paragraph_index": len(document.paragraphs) - 1,
        "backup_path": backup_result["backup_path"],
    }


def insert_before(path: str, index: int, text_value: str, style: str | None = None) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    if index < 0 or index >= len(document.paragraphs):
        raise DctlError("INVALID_SELECTOR", f"Paragraph index {index} is out of range.")
    backup_result = backup(str(doc_path))
    paragraph = document.paragraphs[index].insert_paragraph_before(text_value)
    if style:
        paragraph.style = style
    document.save(doc_path)
    return {
        "path": str(doc_path),
        "inserted_before": index,
        "text": text_value,
        "backup_path": backup_result["backup_path"],
    }


def set_paragraph(path: str, index: int, text_value: str) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    if index < 0 or index >= len(document.paragraphs):
        raise DctlError("INVALID_SELECTOR", f"Paragraph index {index} is out of range.")
    backup_result = backup(str(doc_path))
    document.paragraphs[index].text = text_value
    document.save(doc_path)
    return {
        "path": str(doc_path),
        "paragraph_index": index,
        "text": text_value,
        "backup_path": backup_result["backup_path"],
    }


def replace(path: str, find_text: str, replace_text: str) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    backup_result = backup(str(doc_path))
    replacements = 0
    fallback_resets = 0

    for paragraph in _iter_paragraphs(document):
        run_replacements = 0
        for run in paragraph.runs:
            if find_text in run.text:
                run.text = run.text.replace(find_text, replace_text)
                run_replacements += 1
        if run_replacements:
            replacements += run_replacements
            continue
        if find_text in paragraph.text:
            paragraph.text = paragraph.text.replace(find_text, replace_text)
            replacements += 1
            fallback_resets += 1

    document.save(doc_path)
    return {
        "path": str(doc_path),
        "find": find_text,
        "replace": replace_text,
        "replacements": replacements,
        "paragraph_reset_fallbacks": fallback_resets,
        "backup_path": backup_result["backup_path"],
    }


def _copy_run_appearance(source_run: Any, target_run: Any, *, bold: bool | None = None) -> None:
    if source_run.style:
        target_run.style = source_run.style
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.italic = source_run.font.italic
    target_run.font.underline = source_run.font.underline
    target_run.font.color.rgb = source_run.font.color.rgb
    target_run.font.all_caps = source_run.font.all_caps
    target_run.font.small_caps = source_run.font.small_caps
    target_run.font.highlight_color = source_run.font.highlight_color
    target_run.font.strike = source_run.font.strike
    target_run.font.subscript = source_run.font.subscript
    target_run.font.superscript = source_run.font.superscript
    target_run.bold = bool(source_run.bold) if bold is None else bold


def _clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if paragraph._p.pPr is not None:
        new_paragraph._p.append(deepcopy(paragraph._p.pPr))
    if paragraph.style:
        new_paragraph.style = paragraph.style
    return new_paragraph


def _write_answer_paragraph(paragraph: Paragraph, answer: str, template_runs: list[Any] | None = None) -> None:
    _clear_paragraph(paragraph)
    run = paragraph.add_run(answer)
    if template_runs:
        source_run = next((candidate for candidate in template_runs if candidate.text.strip()), template_runs[0])
        _copy_run_appearance(source_run, run, bold=False)
    else:
        run.bold = False


def _find_question_index(document: Any, question_text: str, exact: bool = False) -> int:
    normalized_target = _normalize_text(question_text)
    matches: list[int] = []
    for index, paragraph in enumerate(document.paragraphs):
        candidate = _normalize_text(paragraph.text)
        if not candidate or not _is_question_paragraph(paragraph):
            continue
        if exact:
            if candidate == normalized_target:
                matches.append(index)
        elif normalized_target in candidate:
            matches.append(index)
    if not matches:
        raise DctlError(
            "ELEMENT_NOT_FOUND",
            f"No question matching '{question_text}' was found.",
            suggestion="Run `dctl docx worksheet-map` to inspect detected questions.",
        )
    if len(matches) > 1:
        raise DctlError(
            "MULTIPLE_MATCHES",
            f"Question selector '{question_text}' matched multiple questions.",
            details={"matches": matches},
        )
    return matches[0]


def _neighbor_answer_paragraph(document: Any, question_index: int) -> Paragraph | None:
    return _next_answer_block(document, document.paragraphs[question_index])


def answer_question(path: str, question: str, answer: str, exact: bool = False) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    question_index = _find_question_index(document, question, exact=exact)
    question_paragraph = document.paragraphs[question_index]
    existing_answer = _neighbor_answer_paragraph(document, question_index)
    backup_result = backup(str(doc_path))

    if existing_answer is not None:
        _write_answer_paragraph(existing_answer, answer, template_runs=list(existing_answer.runs) or list(question_paragraph.runs))
        answer_index = question_index + 1
        action = "replaced"
    else:
        inserted = _insert_paragraph_after(question_paragraph)
        template_runs = list(question_paragraph.runs)
        _write_answer_paragraph(inserted, answer, template_runs=template_runs)
        answer_index = question_index + 1
        action = "inserted"

    document.save(doc_path)
    return {
        "path": str(doc_path),
        "question_index": question_index,
        "answer_index": answer_index,
        "question": question_paragraph.text,
        "answer": answer,
        "action": action,
        "backup_path": backup_result["backup_path"],
    }


def answer_all(path: str, answers_json: str, exact: bool = False) -> dict[str, Any]:
    payload = _parse_json_input(answers_json)
    if isinstance(payload, dict):
        items = [{"question": question, "answer": answer} for question, answer in payload.items()]
    elif isinstance(payload, list):
        items = payload
    else:
        raise DctlError("INVALID_SELECTOR", "Answer payload must be a JSON object or array.")

    results = []
    for item in items:
        if not isinstance(item, dict) or "question" not in item or "answer" not in item:
            raise DctlError("INVALID_SELECTOR", "Each answer item must include `question` and `answer`.")
        results.append(answer_question(path, str(item["question"]), str(item["answer"]), exact=exact))
    return {"path": str(_path(path)), "items": results}


def _match_table_selector(tables: list[dict[str, Any]], selector: str) -> int:
    selector = selector.strip()
    if selector.isdigit():
        index = int(selector)
        if 0 <= index < len(tables):
            return index
        raise DctlError("INVALID_SELECTOR", f"Table index {selector} is out of range.")
    normalized = _normalize_text(selector)
    matches = []
    for table in tables:
        title = _normalize_text(table.get("title") or "")
        if normalized and normalized in title:
            matches.append(table["index"])
    if not matches:
        raise DctlError(
            "ELEMENT_NOT_FOUND",
            f"No table matching '{selector}' was found.",
            suggestion="Run `dctl docx worksheet-map` to inspect table titles and indexes.",
        )
    if len(matches) > 1:
        raise DctlError(
            "MULTIPLE_MATCHES",
            f"Table selector '{selector}' matched multiple tables.",
            details={"matches": matches},
        )
    return matches[0]


def _match_text_in_list(values: list[str], selector: str, label: str) -> int:
    normalized = _normalize_text(selector)
    exact_matches = [index for index, value in enumerate(values) if _normalize_text(value) == normalized]
    if len(exact_matches) == 1:
        return exact_matches[0]
    partial_matches = [index for index, value in enumerate(values) if normalized and normalized in _normalize_text(value)]
    matches = exact_matches or partial_matches
    if not matches:
        raise DctlError("ELEMENT_NOT_FOUND", f"No {label} matching '{selector}' was found.")
    if len(matches) > 1:
        raise DctlError("MULTIPLE_MATCHES", f"{label.title()} selector '{selector}' matched multiple values.", details={label: matches})
    return matches[0]


def _write_cell_text(cell: _Cell, value: str) -> None:
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
        _write_answer_paragraph(paragraph, value)
        return
    first = cell.paragraphs[0]
    template_runs = list(first.runs)
    _write_answer_paragraph(first, value, template_runs=template_runs)
    for paragraph in cell.paragraphs[1:]:
        _clear_paragraph(paragraph)


def fill_table(path: str, table: str, entries_json: str) -> dict[str, Any]:
    doc_path = _path(path)
    document = Document(doc_path)
    worksheet = worksheet_map(str(doc_path))
    table_index = _match_table_selector(worksheet["tables"], table)
    target_table = document.tables[table_index]
    matrix = _table_matrix(target_table)
    if len(matrix) < 2 or len(matrix[0]) < 2:
        raise DctlError("ACTION_NOT_SUPPORTED", "Table is too small to infer headers and row labels.")
    column_headers = matrix[0][1:]
    row_labels = [row[0] for row in matrix[1:]]
    payload = _parse_json_input(entries_json)
    if not isinstance(payload, list):
        raise DctlError("INVALID_SELECTOR", "Table fill payload must be a JSON array.")

    backup_result = backup(str(doc_path))
    results = []
    for item in payload:
        if not isinstance(item, dict) or not {"row_label", "column_label", "value"} <= set(item):
            raise DctlError("INVALID_SELECTOR", "Each table entry must include `row_label`, `column_label`, and `value`.")
        row_index = _match_text_in_list(row_labels, str(item["row_label"]), "row_label") + 1
        column_index = _match_text_in_list(column_headers, str(item["column_label"]), "column_label") + 1
        cell = target_table.cell(row_index, column_index)
        _write_cell_text(cell, str(item["value"]))
        results.append(
            {
                "row_label": item["row_label"],
                "column_label": item["column_label"],
                "row_index": row_index,
                "column_index": column_index,
                "value": str(item["value"]),
            }
        )
    document.save(doc_path)
    return {
        "path": str(doc_path),
        "table_index": table_index,
        "table_title": worksheet["tables"][table_index].get("title"),
        "items": results,
        "backup_path": backup_result["backup_path"],
    }


def diff(path: str, against: str) -> dict[str, Any]:
    current_path = _path(path)
    against_path = _path(against)
    current = read(str(current_path), include_empty=True)["paragraphs"]
    baseline = read(str(against_path), include_empty=True)["paragraphs"]
    diff_lines = list(
        difflib.unified_diff(
            baseline,
            current,
            fromfile=str(against_path),
            tofile=str(current_path),
            lineterm="",
        )
    )
    return {"path": str(current_path), "against": str(against_path), "lines": diff_lines}
