from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document

from dctl.adapters import docx_files


class DocxAdapterTests(unittest.TestCase):
    def test_append_and_replace(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            document = Document()
            document.add_paragraph("hello world")
            document.save(path)

            append_result = docx_files.append(str(path), "second paragraph")
            replace_result = docx_files.replace(str(path), "world", "planet")
            read_result = docx_files.read(str(path))

            self.assertTrue(Path(append_result["backup_path"]).exists())
            self.assertTrue(Path(replace_result["backup_path"]).exists())
            self.assertEqual(replace_result["replacements"], 1)
            self.assertIn("hello planet", read_result["paragraphs"])
            self.assertIn("second paragraph", read_result["paragraphs"])

    def test_insert_before_and_set_paragraph(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            document = Document()
            document.add_paragraph("alpha")
            document.add_paragraph("omega")
            document.save(path)

            docx_files.insert_before(str(path), 1, "middle")
            docx_files.set_paragraph(str(path), 0, "start")
            result = docx_files.paragraphs(str(path))

            self.assertEqual(result["items"][0]["text"], "start")
            self.assertEqual(result["items"][1]["text"], "middle")
            self.assertEqual(result["items"][2]["text"], "omega")

    def test_worksheet_map_and_answer_question_preserve_nonbold_answer(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "worksheet.docx"
            document = Document()
            question = document.add_paragraph()
            run = question.add_run("1. What is 2 + 2?")
            run.bold = True
            document.save(path)

            worksheet = docx_files.worksheet_map(str(path))
            self.assertEqual(len(worksheet["questions"]), 1)

            result = docx_files.answer_question(str(path), "What is 2 + 2?", "4")
            self.assertEqual(result["action"], "inserted")
            self.assertTrue(Path(result["backup_path"]).exists())

            updated = Document(path)
            self.assertEqual(updated.paragraphs[0].text, "1. What is 2 + 2?")
            self.assertEqual(updated.paragraphs[1].text, "4")
            self.assertTrue(updated.paragraphs[0].runs[0].bold)
            self.assertFalse(bool(updated.paragraphs[1].runs[0].bold))

    def test_answer_question_replaces_existing_answer(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "worksheet.docx"
            document = Document()
            document.add_paragraph("2. Capital of France?")
            document.add_paragraph("London")
            document.save(path)

            result = docx_files.answer_question(str(path), "Capital of France", "Paris")
            self.assertEqual(result["action"], "replaced")

            updated = Document(path)
            self.assertEqual(updated.paragraphs[1].text, "Paris")
            self.assertEqual(len(updated.paragraphs), 2)

    def test_fill_table_by_title_and_labels(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "worksheet.docx"
            document = Document()
            document.add_paragraph("Vocabulary Table")
            table = document.add_table(rows=3, cols=3)
            table.cell(0, 0).text = ""
            table.cell(0, 1).text = "Definition"
            table.cell(0, 2).text = "Example"
            table.cell(1, 0).text = "Mitochondria"
            table.cell(2, 0).text = "Chlorophyll"
            document.save(path)

            result = docx_files.fill_table(
                str(path),
                "Vocabulary",
                '[{"row_label":"Mitochondria","column_label":"Definition","value":"Powerhouse of the cell"}]',
            )
            self.assertEqual(result["table_title"], "Vocabulary Table")
            updated = Document(path)
            self.assertEqual(updated.tables[0].cell(1, 1).text, "Powerhouse of the cell")


if __name__ == "__main__":
    unittest.main()
