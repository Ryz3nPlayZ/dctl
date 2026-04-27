from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_dctl_docs_detailed():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('dctl: The Ultimate Desktop Control Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('A comprehensive technical guide for AI Agent Integration')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'dctl (Desktop Control) is a next-generation automation framework built specifically for the LLM agent era. '
        'Unlike traditional automation tools (Selenium, AutoGui) which focus on testing or simple macros, '
        'dctl provides the semantic "sight" and high-precision "hands" an agent needs to reason about and interact with '
        'complex, non-standardized desktop environments.'
    )

    # Core Philosophy
    doc.add_heading('2. Core Philosophy', level=1)
    doc.add_paragraph(
        'The framework operates on three fundamental pillars:'
    )
    p = doc.add_paragraph('Semantic Sight: ', style='List Bullet')
    p.add_run('Prioritizing accessibility trees (UIA, AX, AT-SPI) over pixel-matching.').italic = True
    
    p = doc.add_paragraph('Native Precision: ', style='List Bullet')
    p.add_run('Direct communication with application protocols (CDP for browsers, UNO for LibreOffice).').italic = True
    
    p = doc.add_paragraph('Platform Agnosticism: ', style='List Bullet')
    p.add_run('A single tool-schema for an agent, regardless of whether it is running on a Mac, Windows, or Linux machine.').italic = True

    # Detailed Command Reference
    doc.add_heading('3. Command Reference', level=1)
    
    doc.add_heading('3.1 System Discovery', level=2)
    doc.add_paragraph('dctl list-windows: Returns a structured JSON list of all open windows, including their PIDs and bounds.')
    doc.add_paragraph('dctl doctor: Performs a recursive health check of all required native dependencies (e.g., xdotool, atspi-bus).')

    doc.add_heading('3.2 Browser Control', level=2)
    doc.add_paragraph(
        'The browser module is the most advanced component of dctl. It supports "Session Management" which allows '
        'agents to maintain separate browser profiles (with cookies and history) for different tasks.'
    )
    doc.add_paragraph('dctl browser start --session research: Starts a dedicated browser profile.')
    doc.add_paragraph('dctl browser open https://arxiv.org: Navigates to a URL in the active tab.')

    # Office Precision
    doc.add_heading('4. Office Precision', level=1)
    doc.add_paragraph(
        'One of dctl\'s unique advantages is its ability to edit documents headlessly. This is crucial for '
        '"Background Agents" that need to generate reports without opening a visible window.'
    )
    
    # Table of API endpoints
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Module'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Capability'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Browser'
    row_cells[1].text = 'CDP'
    row_cells[2].text = 'DOM/AX Access'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Word'
    row_cells[1].text = 'python-docx'
    row_cells[2].text = 'Headless Edit'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Writer'
    row_cells[1].text = 'UNO'
    row_cells[2].text = 'Live Control'

    # Optimization Guide
    doc.add_heading('5. Optimization & Performance', level=1)
    doc.add_paragraph(
        'To minimize compute waste and token consumption, agents should use filtered queries. '
        'Instead of fetching the entire UI tree, use the --role and --name flags to narrow the search.'
    )

    doc.save('dctl_documentation_detailed.docx')
    print('Created dctl_documentation_detailed.docx')

if __name__ == "__main__":
    create_dctl_docs_detailed()
