from docx import Document
from docx.shared import Pt

def create_dctl_docs():
    doc = Document()
    doc.add_heading('dctl: Desktop Control CLI & Agent Framework', 0)

    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph('dctl is a cross-platform command-line tool and Python framework designed to give AI agents precision control over desktop environments. It prioritizes semantic interaction (via Accessibility APIs and CDP) over raw coordinate clicking.')

    # Key Features
    doc.add_heading('Key Features', level=1)
    features = [
        ('Universal UI Tree', 'Extract semantic accessibility trees on Windows (UIA), macOS (AX), and Linux (AT-SPI).'),
        ('Deep Browser Integration', 'Direct CDP control over Chromium-based browsers (Brave, Chrome, Edge).'),
        ('Semantic Office Adapters', 'High-level manipulation of Word (.docx) and Excel (.xlsx) files without UI clicking.'),
        ('LibreOffice UNO Bridge', 'Real-time live control of LibreOffice applications.')
    ]
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)

    # Architecture
    doc.add_heading('Architecture', level=1)
    doc.add_paragraph('The DesktopManager acts as a unified facade, routing commands to platform-specific backends. It handles lazy initialization of providers for capture, input, and windowing.')

    # CLI Reference
    doc.add_heading('CLI Quick Start', level=1)
    doc.add_paragraph('dctl tree --app Brave', style='No Spacing')
    doc.add_paragraph('dctl browser open https://google.com', style='No Spacing')
    doc.add_paragraph('dctl docx append docs.docx "New Chapter"', style='No Spacing')

    doc.save('dctl_documentation.docx')
    print('Created dctl_documentation.docx')

if __name__ == "__main__":
    create_dctl_docs()
